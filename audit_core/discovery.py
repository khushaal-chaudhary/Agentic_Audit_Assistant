from __future__ import annotations

import csv
import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from .models import EvidenceRef, IngestionCoverage, IngestionDocument, IngestionRole
from .parsers import PdfExtraction, decode_text, normalize_text, read_pdf_passages


@dataclass(frozen=True)
class RoleSpec:
    role: str
    fields: dict[str, tuple[str, ...]]


def _aliases(canonical: str, *values: str) -> tuple[str, ...]:
    return (canonical, *values)


ROLE_SPECS = (
    RoleSpec(
        "general_ledger",
        {
            "SACHKONTONUMMER": _aliases("SACHKONTONUMMER", "GL_ACCOUNT", "ACCOUNT_NUMBER"),
            "BUCHUNGSNUMMER": _aliases("BUCHUNGSNUMMER", "DOCUMENT_NUMBER", "ENTRY_NUMBER"),
            "BUCHUNGSDATUM": _aliases("BUCHUNGSDATUM", "POSTING_DATE", "BOOKING_DATE"),
            "BUCHUNGSBETRAG": _aliases("BUCHUNGSBETRAG", "POSTING_AMOUNT", "AMOUNT"),
            "BUCHUNGSTEXT": _aliases("BUCHUNGSTEXT", "POSTING_TEXT", "DESCRIPTION"),
            "BENUTZERKENNUNG": _aliases("BENUTZERKENNUNG", "USER_ID", "POSTED_BY"),
        },
    ),
    RoleSpec(
        "manual_journal_ledger",
        {
            "SACHKONTONUMMER": _aliases("SACHKONTONUMMER", "GL_ACCOUNT", "ACCOUNT_NUMBER"),
            "BUCHUNGSNUMMER": _aliases("BUCHUNGSNUMMER", "DOCUMENT_NUMBER", "ENTRY_NUMBER"),
            "BUCHUNGSDATUM": _aliases("BUCHUNGSDATUM", "POSTING_DATE", "BOOKING_DATE"),
            "BUCHUNGSBETRAG": _aliases("BUCHUNGSBETRAG", "POSTING_AMOUNT", "AMOUNT"),
            "BUCHUNGSTEXT": _aliases("BUCHUNGSTEXT", "POSTING_TEXT", "DESCRIPTION"),
            "BENUTZERKENNUNG": _aliases("BENUTZERKENNUNG", "USER_ID", "POSTED_BY"),
            "ERFASSUNGSNUMMER": _aliases(
                "ERFASSUNGSNUMMER", "CAPTURE_NUMBER", "JOURNAL_ID"
            ),
            "PERIODENZUGEHÖRIGKEIT": _aliases(
                "PERIODENZUGEHÖRIGKEIT", "SOURCE_TYPE", "JOURNAL_ORIGIN"
            ),
        },
    ),
    RoleSpec(
        "vendor_postings",
        {
            "LIEFERANTENKONTONUMMER": _aliases(
                "LIEFERANTENKONTONUMMER", "VENDOR_ID", "VENDOR_ACCOUNT", "SUPPLIER_ID"
            ),
            "BUCHUNGSNUMMER": _aliases("BUCHUNGSNUMMER", "DOCUMENT_NUMBER", "ENTRY_NUMBER"),
            "BUCHUNGSDATUM": _aliases("BUCHUNGSDATUM", "POSTING_DATE", "BOOKING_DATE"),
            "BUCHUNGSTEXT": _aliases("BUCHUNGSTEXT", "POSTING_TEXT", "DESCRIPTION"),
            "BUCHUNGSBETRAG": _aliases("BUCHUNGSBETRAG", "POSTING_AMOUNT", "AMOUNT"),
        },
    ),
    RoleSpec(
        "asset_master",
        {
            "ANLAGENNUMMER": _aliases("ANLAGENNUMMER", "ASSET_ID", "ASSET_NUMBER"),
            "ANLAGENBEZEICHNUNG": _aliases(
                "ANLAGENBEZEICHNUNG", "ASSET_DESCRIPTION", "DESCRIPTION"
            ),
            "ANLAGENGRUPPE": _aliases("ANLAGENGRUPPE", "ASSET_GROUP"),
        },
    ),
    RoleSpec(
        "asset_postings",
        {
            "ANLAGENNUMMER": _aliases("ANLAGENNUMMER", "ASSET_ID", "ASSET_NUMBER"),
            "WERTSTELLUNG": _aliases("WERTSTELLUNG", "VALUE_DATE", "POSTING_DATE"),
            "BELEGNUMMER": _aliases("BELEGNUMMER", "DOCUMENT_NUMBER"),
            "BUCHUNGSBETRAG": _aliases("BUCHUNGSBETRAG", "POSTING_AMOUNT", "AMOUNT"),
            "BUCHUNGSART": _aliases("BUCHUNGSART", "POSTING_TYPE", "TRANSACTION_TYPE"),
        },
    ),
    RoleSpec(
        "goods_receipts",
        {
            "WARENEINGANG_NR": _aliases(
                "WARENEINGANG_NR", "GOODS_RECEIPT_NUMBER", "RECEIPT_NUMBER"
            ),
            "WARENEINGANG_DATUM": _aliases(
                "WARENEINGANG_DATUM", "GOODS_RECEIPT_DATE", "RECEIPT_DATE"
            ),
            "KREDITOR": _aliases("KREDITOR", "VENDOR_ID", "VENDOR"),
            "BETRAG_EUR": _aliases("BETRAG_EUR", "AMOUNT_EUR", "AMOUNT"),
            "BEMERKUNG": _aliases("BEMERKUNG", "REMARK", "STATUS"),
        },
    ),
    RoleSpec(
        "vendor_changes",
        {
            "DATUM": _aliases("DATUM", "CHANGE_DATE", "DATE"),
            "ART": _aliases("ART", "CHANGE_TYPE", "TYPE"),
            "KONTO": _aliases("KONTO", "ACCOUNT", "VENDOR_ID"),
            "NAME": _aliases("NAME", "VENDOR_NAME"),
            "FELD": _aliases("FELD", "FIELD", "CHANGED_FIELD"),
            "GEAENDERT_VON": _aliases("GEAENDERT_VON", "CHANGED_BY"),
            "GENEHMIGT_VON": _aliases("GENEHMIGT_VON", "APPROVED_BY"),
        },
    ),
    RoleSpec(
        "permissions",
        {
            "Benutzer": _aliases("Benutzer", "USER", "USER_ID"),
            "Buchen": _aliases("Buchen", "POST_ENTRIES", "POSTING"),
            "Zahlungslauf": _aliases("Zahlungslauf", "PAYMENT_RUN"),
            "Stammdaten/Kreditor anlegen": _aliases(
                "Stammdaten/Kreditor anlegen", "CREATE_VENDOR", "VENDOR_CREATE"
            ),
        },
    ),
    RoleSpec(
        "future_vendor_invoices",
        {
            "RECHNUNGSNUMMER": _aliases("RECHNUNGSNUMMER", "INVOICE_NUMBER"),
            "KREDITOR": _aliases("KREDITOR", "VENDOR_ID", "VENDOR"),
            "FAKTURADATUM": _aliases("FAKTURADATUM", "INVOICE_DATE"),
            "LEISTUNGSDATUM": _aliases("LEISTUNGSDATUM", "SERVICE_DATE"),
            "BETRAG_EUR": _aliases("BETRAG_EUR", "AMOUNT_EUR", "AMOUNT"),
        },
    ),
    RoleSpec(
        "journal_approvals",
        {
            "ERFASSUNGSNUMMER": _aliases(
                "ERFASSUNGSNUMMER", "CAPTURE_NUMBER", "JOURNAL_ID"
            ),
            "JOURNALNAME": _aliases("JOURNALNAME", "JOURNAL_NAME"),
            "ANZAHL_ZEILEN": _aliases("ANZAHL_ZEILEN", "LINE_COUNT"),
            "SUMME_ABS_EUR": _aliases(
                "SUMME_ABS_EUR", "ABSOLUTE_AMOUNT_EUR", "ABSOLUTE_JOURNAL_AMOUNT"
            ),
            "ERSTELLER": _aliases("ERSTELLER", "CREATOR", "CREATED_BY"),
            "ERFASST_AM": _aliases("ERFASST_AM", "CREATED_DATE"),
            "ERFASST_UM": _aliases("ERFASST_UM", "CREATED_TIME"),
            "FREIGEBER": _aliases("FREIGEBER", "APPROVER", "APPROVED_BY"),
            "FREIGABESTATUS": _aliases("FREIGABESTATUS", "APPROVAL_STATUS"),
        },
    ),
)

REQUIRED_ROLES = tuple(spec.role for spec in ROLE_SPECS) + (
    "payment_policy",
    "jet_policy",
    "export_manifest",
    "it_completeness_confirmation",
)


def _token(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", normalize_text(value))


def _header_map(headers: list[str], spec: RoleSpec) -> dict[str, str] | None:
    actual_by_token = {_token(header): header for header in headers if header}
    mapping: dict[str, str] = {}
    for canonical, aliases in spec.fields.items():
        actual = next(
            (
                actual_by_token[_token(alias)]
                for alias in aliases
                if _token(alias) in actual_by_token
            ),
            None,
        )
        if actual is None:
            return None
        mapping[canonical] = actual
    return mapping


def _role_maps(headers: list[str]) -> dict[str, dict[str, str]]:
    return {
        spec.role: mapping
        for spec in ROLE_SPECS
        if (mapping := _header_map(headers, spec)) is not None
    }


def _gdpdu_tables(index_path: Path) -> list[tuple[Path, list[str]]]:
    tree = ET.parse(index_path)
    tables: list[tuple[Path, list[str]]] = []
    for table in tree.iter():
        if table.tag.split("}")[-1] != "Table":
            continue
        url = next(
            (child.text for child in table if child.tag.split("}")[-1] == "URL"),
            None,
        )
        if not url:
            continue
        headers = [
            child.text or ""
            for column in table.iter()
            if column.tag.split("}")[-1] == "VariableColumn"
            for child in column
            if child.tag.split("}")[-1] == "Name"
        ]
        tables.append((index_path.parent / url, headers))
    return tables


def _text_headers(path: Path) -> list[str]:
    first_line = decode_text(path).splitlines()[:1]
    if not first_line:
        return []
    return [value.strip() for value in next(csv.reader(first_line, delimiter=";"))]


def _xlsx_profile(path: Path) -> tuple[list[str], dict[str, dict[str, str]]]:
    workbook = load_workbook(path, read_only=True, data_only=False)
    best_headers: list[str] = []
    best_maps: dict[str, dict[str, str]] = {}
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(max_row=40, values_only=True):
            headers = [str(value or "").strip() for value in row]
            maps = _role_maps(headers)
            if len(maps) > len(best_maps):
                best_headers, best_maps = headers, maps
    workbook.close()
    return best_headers, best_maps


def _policy_match(path: Path) -> bool:
    document = Document(path)
    text = normalize_text(" ".join(paragraph.text for paragraph in document.paragraphs))
    return (
        ("zahlungsfreig" in text or "payment approval" in text)
        and ("vier-augen" in text or "dual approval" in text or "four eyes" in text)
    )


def _jet_policy_match(path: Path) -> bool:
    document = Document(path)
    text = normalize_text(" ".join(paragraph.text for paragraph in document.paragraphs))
    return (
        ("journal entry testing" in text or "nichtaufgriffsgrenze jet" in text)
        and ("nichtaufgriffsgrenze" in text or "clearly trivial threshold" in text)
    )


def pdf_role_matches(passages: list[EvidenceRef]) -> list[str]:
    """Return conservative, descriptive PDF roles from native text only."""
    text = normalize_text("\n".join(passage.excerpt for passage in passages))
    if not text:
        return []
    if (
        ("exportprotokoll" in text or "export log" in text)
        and "sha-256" in text
        and ("datensatze" in text or "records" in text)
    ):
        return ["export_manifest"]
    if (
        ("vollstandigkeit" in text or "completeness" in text)
        and ("unveranderbarkeit" in text or "immutability" in text)
        and "journal" in text
    ):
        return ["it_completeness_confirmation"]
    if (
        (
            ("bilanz" in text or "balance sheet" in text)
            and (
                "gewinn- und verlustrechnung" in text
                or "income statement" in text
                or "profit and loss" in text
            )
        )
        or (
            "summe aktiva" in text
            and "summe passiva" in text
            and ("jahresabschluss" in text or "financial statements" in text)
        )
    ):
        return ["financial_statements"]
    return []


def discover_dossier(root: Path) -> IngestionCoverage:
    root = root.resolve()
    inspected: dict[Path, tuple[str, list[str], dict[str, dict[str, str]]]] = {}
    pdf_extractions: dict[Path, PdfExtraction] = {}
    for index_path in root.rglob("index.xml"):
        try:
            for path, headers in _gdpdu_tables(index_path):
                if path.is_file():
                    inspected[path.resolve()] = ("gdpdu", headers, _role_maps(headers))
        except (ET.ParseError, OSError):
            continue

    for path in sorted(item for item in root.rglob("*") if item.is_file()):
        resolved = path.resolve()
        if resolved in inspected:
            continue
        suffix = path.suffix.casefold()
        if suffix in {".csv", ".txt"}:
            headers = _text_headers(path)
            inspected[resolved] = ("delimited", headers, _role_maps(headers))
        elif suffix == ".xlsx":
            headers, maps = _xlsx_profile(path)
            inspected[resolved] = ("xlsx", headers, maps)
        elif suffix == ".docx":
            maps = {}
            if _policy_match(path):
                maps["payment_policy"] = {}
            if _jet_policy_match(path):
                maps["jet_policy"] = {}
            inspected[resolved] = ("docx", [], maps)
        elif suffix == ".pdf":
            extraction = read_pdf_passages(path, root)
            pdf_extractions[resolved] = extraction
            maps = {role: {} for role in pdf_role_matches(extraction.passages)}
            inspected[resolved] = ("pdf", [], maps)
        else:
            inspected[resolved] = (suffix.lstrip(".") or "file", [], {})

    role_candidates: dict[str, list[tuple[Path, dict[str, str]]]] = {
        role: [] for role in REQUIRED_ROLES
    }
    for path, (_, _, maps) in inspected.items():
        for role, mapping in maps.items():
            if role in role_candidates:
                role_candidates[role].append((path, mapping))

    roles: list[IngestionRole] = []
    for role in REQUIRED_ROLES:
        candidates = role_candidates.get(role, [])
        if len(candidates) == 1:
            path, mapping = candidates[0]
            roles.append(
                IngestionRole(
                    role=role,
                    status="resolved",
                    document=path.relative_to(root).as_posix(),
                    header_map=mapping,
                    reason="Unique schema/content match",
                )
            )
        elif candidates:
            roles.append(
                IngestionRole(
                    role=role,
                    status="ambiguous",
                    reason=(
                        "Multiple schema/content matches: "
                        + ", ".join(path.relative_to(root).as_posix() for path, _ in candidates)
                    ),
                )
            )
        else:
            roles.append(
                IngestionRole(
                    role=role,
                    status="missing",
                    reason="No document matched the required schema/content signature",
                )
            )

    ambiguous_paths = {
        path
        for candidates in role_candidates.values()
        if len(candidates) > 1
        for path, _ in candidates
    }
    documents: list[IngestionDocument] = []
    for path, (format_name, _, maps) in sorted(
        inspected.items(), key=lambda item: item[0].as_posix()
    ):
        matches = sorted(maps)
        if path in ambiguous_paths:
            status = "ambiguous"
            reason = "Document cannot be assigned to one unique role"
        elif matches:
            status = "recognized"
            reason = (
                "Native PDF text matched " if format_name == "pdf" else "Schema/content matched "
            ) + ", ".join(matches)
        elif format_name == "pdf" and pdf_extractions[path].status == "unreadable":
            status = "unsupported"
            reason = "No extractable native PDF text; OCR is disabled"
            if pdf_extractions[path].error:
                reason += ". " + pdf_extractions[path].error
        elif format_name == "pdf":
            status = "unclassified"
            reason = "Native PDF text extracted; no implemented role consumes its content"
        else:
            status = "unclassified"
            reason = "No implemented rule currently consumes this document schema"
        extraction = pdf_extractions.get(path)
        documents.append(
            IngestionDocument(
                document=path.relative_to(root).as_posix(),
                format=format_name,
                status=status,
                role_matches=matches,
                reason=reason,
                extraction_status=extraction.status if extraction else "not_applicable",
                page_count=extraction.page_count if extraction else None,
                extracted_pages=extraction.extracted_pages if extraction else 0,
                passage_count=len(extraction.passages) if extraction else 0,
            )
        )
    source_passages = [
        passage
        for path in sorted(pdf_extractions, key=lambda item: item.as_posix())
        for passage in pdf_extractions[path].passages
    ]
    return IngestionCoverage(
        documents=documents,
        roles=roles,
        source_passages=source_passages,
    )
